#!/usr/bin/env python3
import sys, json, os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

FONT_NAME = "Arial"

def generate_docx(data, client_name, alert_number, output_path):
    doc = Document()
    
    # 1. Setup Standard Margins (Matching A4)
    for section in doc.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        
    doc.styles["Normal"].font.name = FONT_NAME
    doc.styles["Normal"].font.size = Pt(11)
    
    # 2. Cover Logo [Image 1]
    logo_path = os.path.join(os.path.dirname(__file__), "logo.png")
    if os.path.exists(logo_path):
        doc.add_picture(logo_path, width=Inches(2.0))
    else:
        p = doc.add_paragraph("[Image 1] - (Place logo.png in services/reporting/)")
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.add_paragraph("\n")

    # 3. Main Title
    title_text = data.get("title", "Security Advisory")
    cve_text = f"({data.get('cve_ids', '')})"
    
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run(f"{title_text}\n{cve_text}")
    run_title.font.name = FONT_NAME
    run_title.font.size = Pt(14)
    run_title.font.bold = True
    
    doc.add_paragraph("\n")

    # 4. Metadata Block
    meta_table = doc.add_table(rows=4, cols=3)
    meta_data = [
        ("Type", ":", data.get("type", "Security Vulnerability")),
        ("Severity", ":", data.get("severity_detail", data.get("severity", "CRITICAL"))),
        ("Target Platforms", ":", data.get("target_platforms", "")),
        ("Alert Number", ":", alert_number)
    ]
    
    for i, (label, colon, val) in enumerate(meta_data):
        row = meta_table.rows[i]
        row.cells[0].width = Inches(1.5)
        row.cells[1].width = Inches(0.2)
        row.cells[2].width = Inches(4.5)
        
        row.cells[0].paragraphs[0].add_run(label)
        row.cells[1].paragraphs[0].add_run(colon)
        row.cells[2].paragraphs[0].add_run(val)

    # 5. Dual Dates (Matching sample format)
    date_str = data.get("generated_at", datetime.utcnow().strftime("%d %b %Y"))[:10]
    p_dates = doc.add_paragraph(f"\n{date_str}\n{date_str}")
    p_dates.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # 6. Overview Section
    doc.add_heading('Overview', level=2)
    
    overview = data.get("overview_table", [])
    if not overview:
        overview = [{"product": p, "severity": data.get("severity", "HIGH"), "cve_id": data.get("cve_ids", "")} for p in (data.get("affected_products") or ["Affected Product"])[:3]]
    
    ov_table = doc.add_table(rows=1+len(overview), cols=3)
    ov_table.style = 'Table Grid'
    hdr_cells = ov_table.rows[0].cells
    hdr_cells[0].text = 'Product'
    hdr_cells[1].text = 'Severity'
    hdr_cells[2].text = 'Threat'
    
    for i, item in enumerate(overview):
        row_cells = ov_table.rows[i+1].cells
        row_cells[0].text = item.get("product", "")
        row_cells[1].text = item.get("severity", "")
        row_cells[2].text = item.get("cve_id", "")

    doc.add_paragraph("\n")

    # 7. Description Section
    doc.add_heading('Description', level=2)
    doc.add_paragraph(data.get("description", "No description available."))

    # 8. Impact Section
    doc.add_heading('Impact', level=2)
    impacts = data.get("impact", [])
    if isinstance(impacts, list):
        for imp in impacts:
            doc.add_paragraph(imp, style='List Bullet')
    else:
        doc.add_paragraph(str(impacts))

    # 9. Remediations Section
    doc.add_heading('Mitigation / Remediations', level=2)
    doc.add_paragraph(data.get("remediation", "Apply vendor patches."))

    # 10. References Section
    doc.add_heading('References', level=2)
    for ref in data.get("references", []):
        doc.add_paragraph(ref, style='List Bullet')

    # 11. Disclaimer Section
    doc.add_heading('Disclaimer', level=2)
    disclaimer_default = (
        "The information provided here is on an “as is” basis, without warranty of any kind.\n\n"
        "Products that are past their End of General Support dates are not evaluated as part of security advisories. "
        "If your organization has extended support, please use those processes to request assistance. "
        "(if you do not have such support, you need to patch to the latest version or Long-Term Support (LTS) version.)"
    )
    doc.add_paragraph(data.get("disclaimer", disclaimer_default))

    # Save
    doc.save(output_path)
    print(f"Saved: {output_path}")

if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: docx_generator.py <json_data_path> <output_docx_path> [client_name] [alert_number]")
        sys.exit(1)
        
    data = json.load(open(sys.argv[1]))
    output = sys.argv[2]
    client = sys.argv[3] if len(sys.argv) > 3 else "Client"
    alert = sys.argv[4] if len(sys.argv) > 4 else "MITESP-XX"
    generate_docx(data, client, alert, output)
